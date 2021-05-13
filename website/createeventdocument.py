from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE, WD_STYLE
from docx.oxml.table import CT_Row, CT_Tc
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

class CreateEventDocument():
    
    def __init__(self, teacher='', date='', destination='', event_name='', starting_date='', ending_date='', starting_location='', ending_location='', pupil_count='', route='', goal='', event_content='', pupil_list=[], skipped_lessons=''):
        self.document = Document('/Users/Povilas/Desktop/Flask_Web_App/website/template.docx')
        self.document.save('eventform_clone.docx')
        self.document = Document('eventform_clone.docx')
        self.font = self.document.styles['Normal'].font
        self.font.name = 'Times New Roman'
        self.font.size = Pt(12)
        index=0
        paragraphs= self.document.paragraphs
        for line in range(40):
            if index == 4:
                paragraphs[index].clear()
                paragraphs[index].add_run(teacher)
            elif index == 7:
                paragraphs[index].clear()
                paragraphs[index].add_run('Kauno jėzuitų gimnazijos\t\t\t\t\t')
                paragraphs[index].add_run(date)
            elif index==11:
                paragraphs[index].clear()
                paragraphs[index].add_run('DĖL EKSKURSIJOS / IŠVYKOS Į ').bold = True
                paragraphs[index].add_run(destination)
            elif index==13:
                paragraphs[index].clear()
                paragraphs[index].add_run(event_name)
                WD_ALIGN_PARAGRAPH.CENTER
            elif index==18:
                paragraphs[index].clear()
                paragraphs[index].add_run('Išvykimo data ir laikas ')
                paragraphs[index].add_run(starting_date)
            elif index==19:
                paragraphs[index].clear()
                paragraphs[index].add_run('Grįžimo data ir laikas ')
                paragraphs[index].add_run(ending_date)
            elif index==20:
                paragraphs[index].clear()
                paragraphs[index].add_run('Išvykimo vieta ')
                paragraphs[index].add_run(starting_location)
            elif index==21:
                paragraphs[index].clear()
                paragraphs[index].add_run('Grįžimo vieta ')
                paragraphs[index].add_run(ending_location)
            elif index==22:
                paragraphs[index].clear()
                paragraphs[index].add_run('Mokinių skaičius ')
                paragraphs[index].add_run(pupil_count)
            elif index==23:
                paragraphs[index].clear()
                paragraphs[index].add_run('Tikslus maršrutas ')
                paragraphs[index].add_run(route)
            elif index==24:
                paragraphs[index].clear()
                paragraphs[index].add_run('Tikslas, uždaviniai ')
                paragraphs[index].add_run(goal) 
            elif index==25:
                paragraphs[index].clear()
                paragraphs[index].add_run('Programos turinys ')
                paragraphs[index].add_run(event_content)
            elif index==28:
                table = self.document.tables
                                 
                row = table[0].rows[0].cells
                row[0].text = 'Vardas'
                row[1].text = 'Pavardė'
                row[2].text = 'Kodas'

                
                for pupil in pupil_list: 
                    row = table[0].add_row().cells
                    print(pupil.name, pupil.surname, pupil.pupilCode)
                    row[0].text = pupil.name
                    row[1].text = pupil.surname
                    row[2].text = pupil.pupilCode
            elif index==31:
                paragraphs[index].clear()
                paragraphs[index].add_run('10. Prašau mokinius atleisti iš ')
                paragraphs[index].add_run(skipped_lessons+' pamokos(-ų).')
            elif index==34:
                paragraphs[index].clear()
                paragraphs[index].add_run('\t\t       Grupės vadovas ')
                paragraphs[index].add_run(teacher)
            index+=1 

    
    def save(self, path):
        self.document.save(path)
        
if __name__ == "__main__":
    doc = CreateEventDocument(
        teacher='Povilas Kirna',
        date='2021-03-19', 
        destination='Zapyski', 
        event_name='Dviraciu Zygis', 
        starting_date='2021-03-19', 
        ending_date='2021-03-19',
        starting_location='Kaunas',
        ending_location='Kaunas',
        pupil_count='27',
        route='Kaunas-Zapyskis',
        goal='Nuvaziuoti',
        event_content='Kelias',
        pupil_list=[('Povilas', 'Kirna', 'IB2', '')],
        skipped_lessons='1-8',
    )
    user, date = 'PovilasKirna', '2021.03.19' 
    doc.save('website/static/documents/'+user+'_'+date+'.docx')